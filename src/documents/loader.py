import uuid  # Для генерации уникальных идентификаторов
from abc import (
    ABC, # ABC - для создания абстрактных классов
    abstractmethod # abstractmethod - для абстрактных методов
)
from typing import Optional
from pathlib import Path  # Path - для удобной работы с путями файловой системы
from PyPDF2 import PdfReader # Для работы с PDF
from docx import Document as DocxDocument # Для работы с DOCX

from src.logger import logger

from .models import Document

class BaseDocumentLoader(ABC):
    """Абстрактный базовый класс для загрузчиков документов."""

    logger_name: str

    @abstractmethod
    def load(self, source: str) -> Optional[Document]:
        """
        Загружает документ из источника.
        Args:
            source: Путь к файлу или URL
        Returns:
            Document или None в случае ошибки
        """
        raise NotImplementedError("Метод load должен быть реализован в подклассе.")
        # pass  # Абстрактный метод, реализация будет в дочерних классах

    def _logger_info(self, message: str):
        """
        Логирует информационное сообщение с указанием имени класса.

        Args:
            message (str): Сообщение для логирования.
        """
        name = self.__class__.__name__
        logger.info(f"[{name}] {message}")

    def _logger_error(self, message: str):
        """
        Логирует сообщение об ошибке с указанием имени класса.

        Args:
            message (str): Сообщение об ошибке для логирования.
        """
        name = self.__class__.__name__
        logger.error(f"[{name}] {message}")

    def _io_error(self, path: Path, error: str):
        """
        Логирует сообщение об ошибке I/O с указанием имени класса и пути файла.
        """
        name = self.__class__.__name__
        logger.error(f"[{name}] Ошибка I/O при чтении файла {str(path)}: {error}")

    def _encoding_error(self, path: Path, error: str):
        """
        Логирует сообщение об ошибке кодировки с указанием имени класса и пути файла.
        """
        name = self.__class__.__name__
        logger.error(f"[{name}] Ошибка кодировки файла {str(path)}: {error}")

    def _general_error(self, path: Path, error: str):
        """
        Логирует сообщение об общей ошибке с указанием имени класса и пути файла.
        """
        name = self.__class__.__name__
        logger.error(f"[{name}] Ошибка при загрузке файла {str(path)}: {error}")

    def _path_check(self, path: Path) -> bool:
        """
        Проверяет, существует ли указанный путь и является ли он файлом.

        Args:
            path (Path): Путь к файлу.

        Returns:
            bool: True, если путь существует и является файлом, иначе False.
        """
        if not path.exists():
            self._logger_error(f"Файл не найден: {str(path)}")
            return False

        # Проверяем, что это файл, а не директория
        if not path.is_file():
            self._logger_error(f"Указанный путь не является файлом: {str(path)}")
            return False
            
        return True

    def _generate_metadata(self, path: Path) -> dict:
        """
        Генерирует базовые метаданные для документа.
        Args:
            source: Путь к файлу или URL
        Returns:
            Dict с метаданными
        """
        return {
            "source": str(path),  # Полный путь к файлу в виде строки
            "filename": path.name,  # Имя файла с расширением
            "extension": path.suffix.lower(),  # Расширение файла в нижнем регистре
            "size_bytes": path.stat().st_size,  # Размер файла в байтах
        }

# Имплементация класса BaseDocumentLoader для загрузки текстовых файлов
class TextFileLoader(BaseDocumentLoader):
    """Загрузчик текстовых файлов."""

    def __init__(self, encoding: str = "utf-8"):
        """
        Инициализация загрузчика.
        Args:
            encoding: Кодировка файла
        """
        self.encoding = encoding

    def load(self, source: str) -> Optional[Document]:
        try:
            path = Path(source)  # Преобразуем строку пути в объект Path
            
            # Проверяем, что файл существует и является файлом
            if not self._path_check(path):
                return None
            
            # Читаем содержимое файла
            self._logger_info(f"Читаем содержимое файла {source}")
            # Открываем файл с указанной кодировкой
            with path.open(encoding = self.encoding) as file:
                content = file.read()

            self._logger_info(f"Чтение {source} завершено")

            return Document(
                content=content,
                metadata=self._generate_metadata(path),
                doc_id=str(uuid.uuid4())
            )

        except OSError as e:
            self._io_error(path, str(e))
            return None

        except (UnicodeDecodeError) as e:
            self._encoding_error(path, str(e))
            return None
            
        except Exception as e:
            self._general_error(path, str(e))
            return None

class PDFLoader(BaseDocumentLoader):
    """Загрузчик PDF файлов."""
    
    def load(self, source: str) -> Optional[Document]:
        try:
            path = Path(source)  # Преобразуем строку пути в объект Path

            # Проверяем, что файл существует и является файлом
            if not self._path_check(path):
                return None

            # Читаем содержимое файла
            self._logger_info(f"Читаем содержимое PDF файла {source}")

            with path.open(mode="rb") as file:
                pdf = PdfReader(file)
                number_of_pages = len(pdf.pages)
                content = ""
                # Извлекаем текст из всех страниц
                for page in pdf.pages:
                    content += page.extract_text() + "\n"

            self._logger_info(f"Чтение {source} завершено")

            # Добавляем специфичные для PDF метаданные
            metadata = self._generate_metadata(path)
            metadata.update({
                "num_pages": len(pdf.pages),
                "type": "pdf"
            })

            return Document(
                content=content,
                metadata=metadata,
                doc_id=str(uuid.uuid4())
            )

        except OSError as e:
            self._io_error(path, str(e))
            return None

        except (UnicodeDecodeError) as e:
            self._encoding_error(path, str(e))
            return None
            
        except Exception as e:
            self._general_error(path, str(e))
            return None

class DocxLoader(BaseDocumentLoader):
    """Загрузчик DOCX файлов."""
    
    def load(self, source: str) -> Optional[Document]:
        try:
            path = Path(source)  # Преобразуем строку пути в объект Path

            # Проверяем, что файл существует и является файлом
            if not self._path_check(path):
                return None

            self._logger_info(f"Читаем содержимое DOCX файла {source}")

            doc = DocxDocument(source) # Открываем DOCX файл
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            # Добавляем специфичные для DOCX метаданные
            metadata = self._generate_metadata(path)
            metadata.update({
                "type": "docx"
            })

            self._logger_info(f"Чтение {source} завершено")

            return Document(
                content=content,
                metadata=metadata,
                doc_id=str(uuid.uuid4())
            )

        except OSError as e:
            self._io_error(path, str(e))
            return None

        except (UnicodeDecodeError) as e:
            self._encoding_error(path, str(e))
            return None
            
        except Exception as e:
            self._general_error(path, str(e))
            return None

# Загрузчик для всех поддерживаемых файлов из директории
class DirectoryLoader():
    """Загрузчик для всех поддерживаемых файлов из директории."""

    def __init__(self):
        """
        Args:
            loader: Загрузчик документов определенного типа
        """
        # Сохраняем загрузчики, которые будем использовать для каждого файла
        self.loaders = {
            '.txt': TextFileLoader(),
            '.pdf': PDFLoader(),
            '.docx': DocxLoader(),
        }

    def _path_check(self, path: Path) -> bool:
        """
        Проверяет, существует ли указанный путь и является ли он файлом.

        Args:
            path (Path): Путь к файлу.

        Returns:
            bool: True, если путь существует и является файлом, иначе False.
        """
        if not path.exists():
            logger.error(f"Файл не найден: {str(path)}")
            return False

        # Проверяем, что это файл, а не директория
        if not path.is_file():
            logger.error(f"Указанный путь не является файлом: {str(path)}")
            return False
            
        return True

    def _load(self, source: str) -> Optional[Document]:
        """
        Загружает документ, автоматически определяя его тип.
        
        Args:
            source: Путь к файлу
            
        Returns:
            Document или None в случае ошибки
        """
        path = Path(source)

        # Проверяем, что файл существует и является файлом
        if not self._path_check(path):
            return None

        # Определяем расширение файла
        extension = path.suffix.lower()

        loader = self.loaders.get(extension, None)

        if loader is None:
            logger.error(
                "Неподдерживаемый тип файла %s. Поддерживаемые типы: %s",
                extension,
                ", ".join(self.loaders.keys())
            )
            return None

        return loader.load(source)

    def load_from_directory(
        self,
        directory: str,
        glob_pattern: str = "*.*"  # По умолчанию загружаем все файлы
    ) -> list[Document]:
        """
        Загружает все поддерживаемые файлы из директории.
        Args:
            directory: Путь к директории
            glob_pattern: Паттерн для фильтрации файлов
        Returns:
            Список загруженных документов
        """
        path = Path(directory)
        documents: list[Document] = [] # Список для хранения загруженных документов

        if not path.is_dir() or not path.exists():
            logger.error("Директория не найдена: %s", directory)
            return documents

        # Перебираем все файлы в директории, соответствующие паттерну
        file_path_list = path.glob(glob_pattern)

        for file_path in file_path_list:
            # Пытаемся загрузить каждый файл
            if doc := self._load(str(file_path)):
                documents.append(doc)

        logger.info("Загружено %s документов из %s", len(documents), directory)
        return documents
